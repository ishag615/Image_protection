"""
Enhanced PII Redactor with Multiple Safeguarding Strategies
Supports blur, pixelate, replacement, FPE, and full encryption
"""

from PIL import Image
import cv2
import numpy as np
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
import logging
from encryption_engine import EncryptionEngine
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt
import io
import re

logger = logging.getLogger(__name__)


class PIIRedactor:
    """
    Advanced PII redaction with multiple safeguarding strategies
    Supports image redaction, document redaction, and text redaction
    """
    
    def __init__(self, encryption_engine: Optional[EncryptionEngine] = None):
        """
        Initialize PII Redactor
        
        Args:
            encryption_engine: Optional EncryptionEngine instance. Creates new one if not provided.
        """
        self.encryption_engine = encryption_engine or EncryptionEngine()
        self.safeguard_methods = {
            'redact': 'Complete blackout redaction',
            'blur': 'Gaussian blur',
            'pixelate': 'Pixelation effect',
            'replace': 'Replace with [REDACTED]',
            'fpe_encrypt': 'Format-preserving encryption',
            'full_encrypt': 'Full field encryption',
            'keep': 'Keep as-is'
        }
        logger.info("✓ PIIRedactor initialized")
    
    # ========== IMAGE REDACTION ==========
    
    def apply_redaction_to_image(self, image_path: str, entities: List[Dict[str, Any]], 
                                 safeguard_selections: Dict[int, str], 
                                 output_path: str) -> str:
        """
        Apply user-selected safeguarding to image
        
        Args:
            image_path: Path to original image
            entities: List of detected entities with locations
            safeguard_selections: Dict mapping entity index to chosen safeguard method
            output_path: Path to save redacted image
            
        Returns:
            Path to redacted image
        """
        try:
            # Load image
            img = Image.open(image_path).convert('RGB')
            img_array = np.array(img)
            img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            
            # Apply safeguards for each entity
            for entity_idx, entity in enumerate(entities):
                method = safeguard_selections.get(entity_idx, 'redact')
                
                # Get entity location
                start = entity.get('start', 0)
                end = entity.get('end', 0)
                
                # Apply the selected safeguard method
                if method == 'redact':
                    img_cv = self._apply_text_redaction(img_cv, entity)
                elif method == 'blur':
                    img_cv = self._apply_text_blur(img_cv, entity)
                elif method == 'pixelate':
                    img_cv = self._apply_text_pixelate(img_cv, entity)
                elif method in ['replace', 'fpe_encrypt', 'full_encrypt']:
                    # These require text replacement, handled at document level
                    img_cv = self._apply_text_redaction(img_cv, entity)
            
            # Save redacted image
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            img_rgb = cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB)
            result_img = Image.fromarray(img_rgb)
            result_img.save(output_path, quality=95)
            
            logger.info(f"Applied redaction to image: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error applying redaction to image: {e}")
            raise
    
    def _apply_text_redaction(self, img_cv: np.ndarray, entity: Dict[str, Any]) -> np.ndarray:
        """Apply complete blackout redaction to entity location in image"""
        try:
            boxes = entity.get('boxes') or []
            if not boxes:
                return img_cv

            for box in boxes:
                x1, y1, x2, y2 = self._clamped_box(img_cv, box, padding=6)
                cv2.rectangle(img_cv, (x1, y1), (x2, y2), (0, 0, 0), -1)
            
            return img_cv
        except Exception as e:
            logger.warning(f"Error applying text redaction: {e}")
            return img_cv
    
    def _apply_text_blur(self, img_cv: np.ndarray, entity: Dict[str, Any]) -> np.ndarray:
        """Apply blur to entity area"""
        try:
            boxes = entity.get('boxes') or []
            if not boxes:
                return img_cv

            for box in boxes:
                x1, y1, x2, y2 = self._clamped_box(img_cv, box, padding=8)
                roi = img_cv[y1:y2, x1:x2]
                if roi.size == 0:
                    continue
                kernel_size = max(21, min(75, max(roi.shape[:2]) // 2))
                kernel_size = kernel_size if kernel_size % 2 == 1 else kernel_size + 1
                img_cv[y1:y2, x1:x2] = cv2.GaussianBlur(roi, (kernel_size, kernel_size), 0)
            
            return img_cv
        except Exception as e:
            logger.warning(f"Error applying blur: {e}")
            return img_cv
    
    def _apply_text_pixelate(self, img_cv: np.ndarray, entity: Dict[str, Any]) -> np.ndarray:
        """Apply pixelation to entity area"""
        try:
            boxes = entity.get('boxes') or []
            if not boxes:
                return img_cv

            for box in boxes:
                x1, y1, x2, y2 = self._clamped_box(img_cv, box, padding=8)
                roi = img_cv[y1:y2, x1:x2]
                if roi.size == 0:
                    continue
                h, w = roi.shape[:2]
                pixel_size = max(6, min(h, w) // 5)
                small_w = max(1, w // pixel_size)
                small_h = max(1, h // pixel_size)
                small = cv2.resize(roi, (small_w, small_h), interpolation=cv2.INTER_LINEAR)
                img_cv[y1:y2, x1:x2] = cv2.resize(small, (w, h), interpolation=cv2.INTER_NEAREST)
            
            return img_cv
        except Exception as e:
            logger.warning(f"Error applying pixelation: {e}")
            return img_cv

    def _clamped_box(self, img_cv: np.ndarray, box: Dict[str, Any], padding: int = 0) -> Tuple[int, int, int, int]:
        """Normalize an OCR box and keep it inside image bounds."""
        h, w = img_cv.shape[:2]
        x = int(box.get('x', box.get('left', 0)))
        y = int(box.get('y', box.get('top', 0)))
        width = int(box.get('width', box.get('w', 0)))
        height = int(box.get('height', box.get('h', 0)))

        x1 = max(0, x - padding)
        y1 = max(0, y - padding)
        x2 = min(w, x + width + padding)
        y2 = min(h, y + height + padding)
        return x1, y1, x2, y2
    
    # ========== DOCUMENT REDACTION ==========
    
    def apply_redaction_to_document(self, doc_path: str, entities: List[Dict[str, Any]], 
                                   safeguard_selections: Dict[int, str], 
                                   output_path: str, file_type: str = 'word') -> str:
        """
        Apply safeguarding to text document
        
        Args:
            doc_path: Path to original document
            entities: List of detected entities
            safeguard_selections: Dict mapping entity index to safeguard method
            output_path: Path to save protected document
            file_type: 'word', 'pdf', or 'image'
            
        Returns:
            Path to protected document
        """
        try:
            if file_type == 'word':
                return self._redact_word_document(doc_path, entities, safeguard_selections, output_path)
            else:
                # For PDFs and images, apply image redaction
                return self.apply_redaction_to_image(doc_path, entities, safeguard_selections, output_path)
        
        except Exception as e:
            logger.error(f"Error applying document redaction: {e}")
            raise
    
    def _redact_word_document(self, doc_path: str, entities: List[Dict[str, Any]], 
                            safeguard_selections: Dict[int, str], 
                            output_path: str) -> str:
        """
        Apply safeguarding to Word document
        Can handle encryption by creating new document with redacted content
        """
        try:
            doc = DocxDocument(doc_path)
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Prepare text content for processing
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            
            # Build replacement map based on safeguard selections
            replacements = {}
            
            for entity_idx, entity in enumerate(entities):
                method = safeguard_selections.get(entity_idx, 'redact')
                original_value = entity.get('raw_value') or entity.get('value', 'Hidden')
                entity_type = entity.get('type', 'UNKNOWN')
                
                if method == 'replace':
                    replacements[original_value] = '[REDACTED]'
                elif method == 'fpe_encrypt':
                    try:
                        encrypted = self.encryption_engine.encrypt_fpe(original_value, entity_type)
                        replacements[original_value] = encrypted
                    except Exception as e:
                        logger.warning(f"FPE encryption failed for {entity_type}, using replacement")
                        replacements[original_value] = '[ENCRYPTED]'
                elif method == 'full_encrypt':
                    try:
                        encrypted = self.encryption_engine.encrypt_full_field(original_value, entity_type)
                        replacements[original_value] = '[ENCRYPTED_' + entity_type[:3].upper() + ']'
                    except Exception as e:
                        logger.warning(f"Full encryption failed, using replacement")
                        replacements[original_value] = '[ENCRYPTED]'
                elif method == 'redact':
                    replacements[original_value] = '[REDACTED]'
            
            # Apply replacements to document
            for paragraph in doc.paragraphs:
                for original, replacement in replacements.items():
                    if original in paragraph.text:
                        # Replace text while preserving some formatting
                        paragraph.text = paragraph.text.replace(original, replacement)
                        
                        # Update run colors to indicate redaction
                        for run in paragraph.runs:
                            if replacement in run.text:
                                run.font.color.rgb = RGBColor(220, 53, 69)  # Red
                                run.font.bold = True

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for original, replacement in replacements.items():
                                if original in paragraph.text:
                                    paragraph.text = paragraph.text.replace(original, replacement)
            
            # Save protected document
            doc.save(output_path)
            logger.info(f"Protected Word document saved to {output_path}")
            
            return output_path
        
        except Exception as e:
            logger.error(f"Error redacting Word document: {e}")
            # Fallback: copy original
            import shutil
            shutil.copy(doc_path, output_path)
            return output_path
    
    # ========== TEXT ENCRYPTION HELPERS ==========
    
    def encrypt_entity_text(self, value: str, entity_type: str, method: str) -> str:
        """
        Encrypt entity text based on selected method
        
        Args:
            value: Original value
            entity_type: Type of entity (e.g., 'US_SSN', 'CREDIT_CARD')
            method: Safeguard method ('fpe_encrypt' or 'full_encrypt')
            
        Returns:
            Encrypted or replaced value
        """
        try:
            if method == 'fpe_encrypt':
                return self.encryption_engine.encrypt_fpe(value, entity_type)
            elif method == 'full_encrypt':
                return self.encryption_engine.encrypt_full_field(value, entity_type)
            elif method == 'replace':
                return '[REDACTED]'
            elif method == 'redact':
                return '[REDACTED]'
            else:
                return value
        except Exception as e:
            logger.warning(f"Encryption failed: {e}, using replacement")
            return '[REDACTED]'
    
    # ========== UTILITY METHODS ==========
    
    def blur_image(self, image_path: str, output_path: str, blur_strength: int = 51) -> str:
        """
        Apply simple blur to entire image (user-facing feature)
        
        Args:
            image_path: Path to image
            output_path: Path to save blurred image
            blur_strength: Blur kernel size (odd number)
            
        Returns:
            Path to blurred image
        """
        try:
            img = Image.open(image_path).convert('RGB')
            img_array = np.array(img)
            img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            
            kernel_size = blur_strength if blur_strength % 2 == 1 else blur_strength + 1
            blurred = cv2.GaussianBlur(img_cv, (kernel_size, kernel_size), 0)
            blurred = cv2.medianBlur(blurred, kernel_size)
            
            img_rgb = cv2.cvtColor(blurred, cv2.COLOR_BGR2RGB)
            result = Image.fromarray(img_rgb)
            
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            result.save(output_path, quality=95)
            
            logger.info(f"Image blurred and saved to {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error blurring image: {e}")
            raise
    
    def pixelate_image(self, image_path: str, output_path: str, pixel_size: int = 20) -> str:
        """
        Apply pixelation to entire image (user-facing feature)
        
        Args:
            image_path: Path to image
            output_path: Path to save pixelated image
            pixel_size: Size of pixelation blocks
            
        Returns:
            Path to pixelated image
        """
        try:
            img = Image.open(image_path).convert('RGB')
            img_array = np.array(img)
            img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            
            h, w = img_cv.shape[:2]
            small = cv2.resize(img_cv, (w // pixel_size, h // pixel_size), interpolation=cv2.INTER_LINEAR)
            pixelated = cv2.resize(small, (w, h), interpolation=cv2.INTER_NEAREST)
            
            img_rgb = cv2.cvtColor(pixelated, cv2.COLOR_BGR2RGB)
            result = Image.fromarray(img_rgb)
            
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            result.save(output_path, quality=95)
            
            logger.info(f"Image pixelated and saved to {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error pixelating image: {e}")
            raise
    
    def get_safeguard_methods(self) -> Dict[str, str]:
        """Get available safeguard methods"""
        return self.safeguard_methods.copy()
