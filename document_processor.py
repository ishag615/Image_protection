import cv2
import numpy as np
from pathlib import Path
from pdf2image import convert_from_path
from docx import Document as DocxDocument

class DocumentProcessor:
    """Process various document formats (images, PDFs, Word docs) for PII detection and redaction"""
    
    def __init__(self):
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.gif'],
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc']
        }
    
    def get_file_type(self, filepath):
        """Determine file type from extension"""
        ext = Path(filepath).suffix.lower()
        
        if ext in self.supported_formats['image']:
            return 'image'
        elif ext in self.supported_formats['pdf']:
            return 'pdf'
        elif ext in self.supported_formats['docx']:
            return 'docx'
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def convert_to_images(self, filepath, file_type):
        """Convert document to list of image paths"""
        output_dir = Path('temp_images')
        output_dir.mkdir(exist_ok=True)
        
        image_paths = []
        
        if file_type == 'image':
            # Already an image, just return its path
            image_paths.append(filepath)
        
        elif file_type == 'pdf':
            # Convert PDF pages to images
            pdf_images = convert_from_path(filepath, dpi=200)
            for idx, page_image in enumerate(pdf_images):
                img_path = output_dir / f'{Path(filepath).stem}_page_{idx}.jpg'
                page_image.save(str(img_path), 'JPEG')
                image_paths.append(str(img_path))
        
        elif file_type == 'docx':
            # Convert DOCX to images via extraction
            doc = DocxDocument(filepath)
            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    # Create image from text (simplified approach)
                    img = self._text_to_image(para.text)
                    img_path = output_dir / f'{Path(filepath).stem}_para_{idx}.jpg'
                    img.save(str(img_path), 'JPEG')
                    image_paths.append(str(img_path))
        
        return image_paths
    
    def redact_image(self, image_path, mask_instructions):
        """Redact sensitive areas from image based on instructions"""
        img = cv2.imread(image_path)
        
        if img is None:
            raise ValueError(f"Could not read image: {image_path}")
        
        # Apply redaction based on mask instructions
        for instruction in mask_instructions:
            # Parse instruction and apply redaction
            img = self._apply_redaction(img, instruction)
        
        # Save redacted image
        output_path = str(Path(image_path).parent / f'redacted_{Path(image_path).name}')
        cv2.imwrite(output_path, img)
        
        return output_path
    
    def _apply_redaction(self, img, instruction):
        """Apply redaction to image based on instruction"""
        # This is a simplified version - in production, you'd parse coordinates
        # from the instruction and apply precise masking
        
        # For now, apply a simple blur/black box approach
        if instruction:
            try:
                # Try to parse coordinates if provided
                # Example: "coordinates: x1,y1,x2,y2"
                if 'coordinates' in instruction.lower():
                    coords = instruction.split(':')[1].strip().split(',')
                    x1, y1, x2, y2 = map(int, coords)
                    # Apply black box redaction
                    img[y1:y2, x1:x2] = [0, 0, 0]
            except:
                # If parsing fails, apply blur to entire image (conservative approach)
                img = cv2.GaussianBlur(img, (99, 99), 0)
        
        return img
    
    def rebuild_document(self, image_paths, file_type, original_filename):
        """Rebuild document from redacted images"""
        output_dir = Path('protected_documents')
        output_dir.mkdir(exist_ok=True)
        
        if file_type == 'image':
            # If single image, just copy the redacted version
            output_path = output_dir / f'protected_{original_filename}'
            if image_paths:
                import shutil
                shutil.copy(image_paths[0], str(output_path))
            return str(output_path)
        
        elif file_type == 'pdf':
            # Convert images back to PDF
            from PIL import Image as PILImage
            images = [PILImage.open(path).convert('RGB') for path in image_paths]
            output_path = output_dir / f'protected_{Path(original_filename).stem}.pdf'
            images[0].save(str(output_path), save_all=True, append_images=images[1:])
            return str(output_path)
        
        elif file_type == 'docx':
            # Create new DOCX from images
            doc = DocxDocument()
            for img_path in image_paths:
                doc.add_paragraph(f"[Redacted content from {Path(img_path).name}]")
            output_path = output_dir / f'protected_{original_filename}'
            doc.save(str(output_path))
            return str(output_path)
        
        return None
    
    def _text_to_image(self, text):
        """Convert text to PIL Image (simplified)"""
        # Create a simple image from text
        from PIL import Image as PILImage, ImageDraw, ImageFont
        
        # Create image
        img = PILImage.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            # Try to use default font
            draw.text((10, 10), text, fill='black')
        except:
            # Fallback to default font
            draw.text((10, 10), text, fill='black')
        
        return img
